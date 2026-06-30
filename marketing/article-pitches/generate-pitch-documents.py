"""
Generate Word pitch packages for DataBloomer trade publication outreach.
Run: python generate-pitch-documents.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).resolve().parent
SITE = "https://www.databloomer.com"
AUTHOR = "Tom [Last Name]"
AUTHOR_TITLE = "Founder, DataBloomer"
AUTHOR_EMAIL = "hello@databloomer.com"

ARTICLE_TITLE = (
    "The Replacement Window Is on the Map: "
    "How Miami Roofers Can Find 13–25 Year Roofs Using Public Data"
)

ARTICLE_SUBTITLE = (
    "A field guide to permit records, aging-inventory math, and why "
    "door-knocking without a map burns fuel in Miami-Dade."
)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_section_break(doc: Document, label: str) -> None:
    doc.add_page_break()
    add_heading(doc, label, level=1)


def add_pitch_letter(
    doc: Document,
    outlet: str,
    editor_line: str,
    custom_open: str,
    custom_angle: str,
) -> None:
    add_heading(doc, f"Editorial Pitch — {outlet}", level=0)
    add_para(doc, f"To: {editor_line}")
    add_para(doc, f"From: {AUTHOR}, {AUTHOR_TITLE}")
    add_para(doc, f"Email: {AUTHOR_EMAIL}")
    add_para(doc, f"Re: Proposed article — “{ARTICLE_TITLE}”")
    doc.add_paragraph()
    add_para(doc, "Dear Editor,")
    add_para(doc, custom_open)
    add_para(doc, custom_angle)
    add_para(
        doc,
        "The piece is written as a practical field guide for sales managers and "
        "canvassing crews—not a product announcement. It walks readers through "
        "Miami-Dade’s public permit and property records, explains the 13–25 year "
        "replacement window, and shows why map-first routing matters in a spread-out "
        "market. I disclose at the end that I built a subscription tool that "
        "automates parts of this workflow; the methodology stands on its own.",
    )
    add_para(
        doc,
        "Suggested length: ~1,900 words. I can supply high-res map screenshots "
        "(anonymized addresses) and a simple one-page sidebar on reading county "
        "permit types if helpful.",
    )
    add_para(doc, "Thank you for your consideration,")
    add_para(doc, f"{AUTHOR}\n{AUTHOR_TITLE}\n{AUTHOR_EMAIL}\n{SITE}")


def add_submission_instructions(doc: Document, outlet: str, instructions: list[str]) -> None:
    add_section_break(doc, f"Submission Instructions — {outlet}")
    add_para(
        doc,
        "Follow these steps before sending the pitch. Verify URLs on the publication "
        "website in case contacts change.",
    )
    for i, step in enumerate(instructions, 1):
        add_para(doc, f"{i}. {step}")


def add_article_body(doc: Document, outlet_key: str) -> None:
    add_section_break(doc, "Proposed Article (ready to edit)")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(ARTICLE_TITLE)
    tr.bold = True
    tr.font.size = Pt(16)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(ARTICLE_SUBTITLE)
    sr.italic = True
    doc.add_paragraph()
    add_para(doc, f"By {AUTHOR} | ~1,900 words | Photos/maps: 2 suggested", italic=True)
    doc.add_paragraph()

    # --- Intro varies slightly ---
    if outlet_key == "frsa":
        intro_extra = (
            "Florida roofers know hurricane cycles and code changes; fewer teams "
            "systematically mine the quiet signal sitting in county GIS: homes whose "
            "last re-roof permit is crossing the 13–25 year band. In Miami-Dade—"
            "one of the densest replacement markets in the state—that signal is "
            "especially loud."
        )
    elif outlet_key == "qualified_remodeler":
        intro_extra = (
            "Replacement contractors outside roofing will recognize the pattern: "
            "a building product hits an age band, permits thin out, and the "
            "neighborhood still looks fine from the street. Roofs are the clearest "
            "example in South Florida, where the early-2000s build cycle is now "
            "delivering a predictable wave of work."
        )
    else:
        intro_extra = (
            "Miami-Dade added tens of thousands of homes between 2001 and 2013. "
            "Those roofs are not failures yet—but they are entering the band where "
            "owners start taking calls, insurers ask questions, and savvy contractors "
            "want to be first on the block."
        )

    add_para(
        doc,
        intro_extra
        + " This article explains how to find those properties using public permit "
        "and property data, and why canvassing without a map is an expensive habit "
        "in a county where ‘20 minutes away’ can mean three different housing eras.",
    )

    add_heading(doc, "Why 13–25 years?", level=2)
    add_para(
        doc,
        "Shingle and tile assemblies in South Florida do not die on a fixed birthday, "
        "but sales teams still need a workable window. Thirteen years is roughly when "
        "original roofs from the last cycle stop looking ‘new’ in owner minds; "
        "seventeen is often cited as a sweet spot for proactive replacement "
        "conversations; twenty-five is a practical upper bound before you are "
        "chasing emergency work instead of planned replacements.",
    )
    add_para(
        doc,
        "The goal is not to knock every 18-year roof. The goal is to rank "
        "neighborhoods where the concentration of aging inventory justifies a crew "
        "week—not a scattershot Saturday.",
    )

    add_heading(doc, "What public data actually gives you", level=2)
    add_para(
        doc,
        "You do not need a proprietary database to start. Miami-Dade publishes "
        "building permits, code enforcement cases, and property appraiser parcels "
        "through county GIS. Useful fields include year built, folio (parcel ID), "
        "issue date of roofing permits (application types for re-roof and repair), "
        "and open structure or minimum-housing violations that may correlate with "
        "roof-related distress.",
    )
    add_bullets(
        doc,
        [
            "Last roofing permit date → proxy for roof age when available.",
            "Year built → fallback when no roof permit exists (common on older stock).",
            "Recent re-roof in the last 3–5 years → exclude; someone already won that job.",
            "Open code cases → secondary list for homes that may already be motivated.",
        ],
    )
    add_para(
        doc,
        "A sales manager can pull a ZIP, filter by year built or last permit decade, "
        "and export a first-pass list in an afternoon. The bottleneck is not access—"
        "it is keeping the list current and making it usable for crews on the street.",
    )

    add_heading(doc, "A one-week workflow (no software required)", level=2)
    add_para(doc, "Monday — Pick two ZIPs your company already services or wants to grow in.")
    add_para(
        doc,
        "Tuesday — In county permit search, pull roofing permits (re-roof / repair types) "
        "and note issue dates for target streets.",
    )
    add_para(
        doc,
        "Wednesday — Cross-check property appraiser records for year built on homes "
        "without roof permits; flag 13–25 year proxies.",
    )
    add_para(
        doc,
        "Thursday — Layer open code enforcement if your team handles distressed properties.",
    )
    add_para(
        doc,
        "Friday — Plot surviving addresses on a map (even a spreadsheet with lat/long "
        "pasted into a free map tool). Draw blocks, not individual doors at random.",
    )
    add_para(
        doc,
        "Teams that run this monthly catch new permits and aging inventory as county "
        "data refreshes. Teams that run it once a quarter still beat crews flying blind.",
    )

    add_heading(doc, "Why canvassing without a map wastes fuel in Miami", level=2)
    add_para(
        doc,
        "Miami-Dade is not one market. Kendall tract homes, Brickell high-rises, and "
        "Homestead’s 2000s subdivisions can sit inside the same ‘Miami’ label on a "
        "lead sheet but behave like different products. Driving between them without "
        "clustering is how canvassers burn half a tank to hit six doors—two of which "
        "had a re-roof permit last year.",
    )
    add_bullets(
        doc,
        [
            "Sprawl: long distances between high-density aging blocks.",
            "Mixed eras on the same arterial road—2004 builds next to 2019 builds.",
            "Condo vs. single-family rules and access differences by submarket.",
            "Repeating streets your competitor already saturated last storm season.",
        ],
    )
    add_para(
        doc,
        "Map-first canvassing means: choose a color-coded block (or tier), park once, "
        "walk a loop, log outcomes, move to the next cluster. Fuel goes to doors with "
        "the highest replacement probability, not to ‘it looked like a nice neighborhood "
        "from the highway.’",
    )

    add_heading(doc, "Scoring beats sorting by street name", level=2)
    add_para(
        doc,
        "Once you have age proxies, add two more dimensions: home value proxy (heated "
        "living area or assessed value when published) and permit confidence (was the "
        "roof date observed or inferred from year built?). A simple 0–100 score lets "
        "reps knock reds and oranges first, yellows when the calendar is light, greens "
        "when you are training someone new.",
    )
    add_para(
        doc,
        "This is the same logic storm restorers use for ‘hot squares,’ applied to "
        "aging inventory instead of hail swaths.",
    )

    add_heading(doc, "Neighborhood momentum—not just individual doors", level=2)
    add_para(
        doc,
        "Advanced teams watch permit velocity by ZIP: where re-roof permits are rising "
        "quarter over quarter while a large aging pool remains. That combination "
        "suggests owners are already saying yes in the area—your pitch is easier when "
        "the block is talking about roofs.",
    )
    add_para(
        doc,
        "Pair aging inventory with permit momentum and you get a deployable forecast: "
        "send canvassers to rising ZIPs before the whole tract turns into tire-kicker "
        "tours.",
    )

    add_heading(doc, "When automation helps (disclosure)", level=2)
    add_para(
        doc,
        "Manual GIS work is fine for a single ZIP. It breaks down when you want "
        "county-wide coverage, weekly refreshes, and a crew-friendly map on phones. "
        "I built DataBloomer (databloomer.com) for Miami-Dade contractors to automate "
        "this pipeline: ingest public permits and parcels, score aging roofs, color-code "
        "Bloom Zones on a map, and surface code enforcement alongside replacement-likely "
        "pins. Subscribers get full addresses and folios; a free preview shows the "
        "method without exposing homeowner data.",
    )
    add_para(
        doc,
        "Whether you use a tool or a spreadsheet, the discipline is the same: public "
        "data → age window → map clusters → measured canvassing. The contractors who "
        "win the next Miami replacement wave will not be the ones with the most door "
        "hangers—they will be the ones who knew which blocks to hit first.",
    )

    add_heading(doc, "Sidebar: Miami-Dade roofing permit types (quick reference)", level=2)
    add_bullets(
        doc,
        [
            "Type 13 — Re-roof / repair (common for age tracking).",
            "Type 18 — Re-roof / repair (hurricane-related category in county data).",
            "Type 20 — Roof recovery.",
            "Always confirm status (issued/finaled) before assuming work completed.",
        ],
    )

    add_heading(doc, "About the author", level=2)
    add_para(
        doc,
        f"{AUTHOR} is {AUTHOR_TITLE}, a Miami-based platform that aggregates "
        f"Miami-Dade County GIS permits, property records, and code enforcement for "
        f"roofing contractors. Contact: {AUTHOR_EMAIL}.",
    )


OUTLETS = [
    {
        "key": "roofing_contractor",
        "filename": "01-Roofing-Contractor-RC-Pitch.docx",
        "outlet": "Roofing Contractor magazine",
        "editor": "Editorial team, Roofing Contractor (RC) — roofingcontractor.com",
        "open": (
            "I am pitching an exclusive field guide for your sales-and-marketing "
            "readers in storm and non-storm markets: how to build a replacement "
            "pipeline from public permit data in Miami-Dade, and why map-based "
            "canvassing beats route lists in sprawling metros."
        ),
        "angle": (
            "RC’s audience balances production and business development; this piece "
            "gives owners a repeatable weekly workflow their canvassing managers can "
            "run without buying another generic lead list."
        ),
        "instructions": [
            "Visit https://www.roofingcontractor.com and open Contact / About Us for current editorial emails.",
            "Identify the editor for Business / Sales or Features (names change—verify on masthead).",
            "Subject line: “Pitch: Miami permit-data canvassing guide (exclusive)”",
            "Attach this Word file OR paste the article section into the email body with a 2-paragraph pitch letter.",
            "Include 2–3 proposed pull quotes (e.g., on map-first canvassing and the 13–25 year window).",
            "Offer anonymized map screenshots; blur addresses if needed.",
            "If no response in 10 business days, follow up once, then try a different editor on the masthead.",
            "Do not mention pricing in the pitch—keep DataBloomer in the article disclosure only.",
        ],
    },
    {
        "key": "roofing_magazine",
        "filename": "02-Roofing-Magazine-NRCA-Publication-Pitch.docx",
        "outlet": "Roofing Magazine (NRCA publication)",
        "editor": "Editor, Roofing Magazine — nrca.net / roofingmagazine.com",
        "open": (
            "For Roofing Magazine readers, I propose a data-driven operations article: "
            "how contractors can use publicly available permit and property records to "
            "find homes in the 13–25 year roof-age window, with Miami-Dade as a worked "
            "example applicable to any county that publishes GIS permits."
        ),
        "angle": (
            "NRCA members increasingly compete on marketing efficiency; this teaches a "
            "defensible, non-spam prospecting method aligned with professional "
            "contracting—not purchased homeowner lists."
        ),
        "instructions": [
            "Go to https://www.nrca.net/membership/roofing-magazine or search “Roofing Magazine editorial guidelines NRCA”.",
            "Pitch to the magazine editor (distinct from NRCA membership marketing).",
            "Note: NRCA membership is not required to pitch, but member angle helps.",
            "Subject: “Pitch: Public permit data for aging-roof prospecting (Miami case study)”",
            "Emphasize educational value for all members, not only Florida.",
            "Add a short “Applying this in your county” closing paragraph if editor requests localization.",
            "Typical lead time: 4–8 weeks; ask about art deadline for maps.",
            "Follow NRCA style: practical, professional tone; avoid superlatives about any one vendor.",
        ],
    },
    {
        "key": "roofing_insights",
        "filename": "03-Roofing-Insights-Pitch.docx",
        "outlet": "Roofing Insights",
        "editor": "Production / editorial team — roofinginsights.com",
        "open": (
            "Roofing Insights reaches owners who watch long-form breakdowns, not "
            "brochures. I would love to adapt this guide as a contributed article "
            "or interview segment: “How Miami crews find 13–25 year roofs without "
            "buying junk lists,” including a live look at county GIS and map-based "
            "canvassing economics (fuel, close rates, block selection)."
        ),
        "angle": (
            "Your audience responds to blunt talk about marketing waste. The hook: "
            "“You're burning fuel because your list isn't a map.” I can demo public "
            "data live and keep brand mentions to a closing disclosure."
        ),
        "instructions": [
            "Visit https://roofinginsights.com and use Contact / Advertise / Media inquiry forms.",
            "Also reach out via their YouTube or LinkedIn presence for faster routing.",
            "Pitch BOTH formats: (A) written guest post, (B) 15–20 min interview with screen share.",
            "Subject: “Story pitch: Map-first canvassing + public permit data (Miami)”",
            "Offer B-roll: map zoom, permit search, before/after route efficiency story.",
            "Prepare 3 talking points on fraud/bad leads industry-wide—aligns with RI tone.",
            "If pitching video, send a 60-second Loom intro + link to databloomer.com/promo.",
            "Disclose founder relationship on air; focus 90% on methodology.",
        ],
    },
    {
        "key": "qualified_remodeler",
        "filename": "04-Qualified-Remodeler-Pitch.docx",
        "outlet": "Qualified Remodeler",
        "editor": "Editor, Qualified Remodeler — qualifiedremodeler.com",
        "open": (
            "Qualified Remodeler readers run replacement businesses where timing "
            "matters—windows, siding, and roofing. I propose a cross-trade piece "
            "framed around aging building components, using Miami-Dade roof permits "
            "as the clearest public-data example of how to find homes entering a "
            "replacement window before the street looks ‘obvious.’"
        ),
        "angle": (
            "Remodeling leaders understand design cycles; this connects roof permit "
            "history to canvassing ROI and fuel costs—relevant to any exterior "
            "replacement contractor, not only Florida roofers."
        ),
        "instructions": [
            "Visit https://www.qualifiedremodeler.com — locate Contact or Editorial submission page.",
            "Subject: “Pitch: Finding replacement-window homes with public permit data”",
            "Frame for remodeling audience: roof as case study, methodology portable to other trades.",
            "Include a short sidebar: “Questions to ask your county GIS department.”",
            "Provide author headshot and 2-sentence bio.",
            "Ask about digital vs. print timing; offer web-only version if faster.",
            "Follow up once after 2 weeks.",
        ],
    },
    {
        "key": "frsa",
        "filename": "05-FRSA-Florida-Roofing-Pitch.docx",
        "outlet": "FRSA (Florida Roofing & Sheet Metal Contractors Association)",
        "editor": "FRSA editorial / communications — floridaroof.com",
        "open": (
            "For FRSA members, I propose a Florida-specific article on mining "
            "Miami-Dade (and by extension other Florida counties’) public permit "
            "data to find 13–25 year roofs, plus a hard look at why unmapped "
            "canvassing wastes fuel across South Florida’s sprawled ZIP codes."
        ),
        "angle": (
            "FRSA’s audience faces hurricane cycles, licensing scrutiny, and "
            "intense local competition. This is a member education piece—how to "
            "prospect professionally using records taxpayers already paid for."
        ),
        "instructions": [
            "Visit https://www.floridaroof.com — find Contact, Magazine, or Newsletter submission info.",
            "CC membership/marketing if editorial email is unclear.",
            "Subject: “Member education pitch: Public permit prospecting in Miami-Dade”",
            "Mention willingness to present a 30-minute webinar for members (strong FRSA angle).",
            "Offer member benefit: discount code for FRSA members if leadership prefers (negotiate separately).",
            "Emphasize Florida Building Code context and hurricane re-roof permit types.",
            "Ask if they prefer Florida Roofing Contractor magazine vs. e-newsletter format.",
            "Follow FRSA conference calendar—pitch 90 days before annual convention for program inclusion.",
        ],
    },
    {
        "key": "nrca",
        "filename": "06-NRCA-Association-Resources-Pitch.docx",
        "outlet": "NRCA (association — ProSpeaker, blog, member resources)",
        "editor": "NRCA membership / ProContractor / education team — nrca.net",
        "open": (
            "Separate from Roofing Magazine editorial, I am reaching NRCA’s membership "
            "education channel with a resource members can use tomorrow: a guide to "
            "aging-roof prospecting from public GIS, illustrated with Miami-Dade, "
            "applicable to any member willing to query their county permit portal."
        ),
        "angle": (
            "NRCA positions members as professionals. This content supports ethical "
            "prospecting and efficient canvassing—good fit for ProContractor e-news, "
            "member blog, or a webinar slot."
        ),
        "instructions": [
            "Do NOT confuse with Roofing Magazine pitch—route to membership education / ProContractor.",
            "Start at https://www.nrca.net — search Contact, “submit article,” or education department.",
            "Subject: “Member resource pitch: County permit data for aging-roof leads”",
            "Offer formats: (1) 1,500-word article, (2) webinar, (3) downloadable checklist PDF.",
            "If pitching webinar, propose title: “Map-first canvassing with public data.”",
            "NRCA may require speaker/vendor disclosure—provide transparent DataBloomer disclosure.",
            "Ask about Member Value Partner program if article is declined (sponsorship path).",
            "Timeline: association content often plans 2–3 months ahead.",
        ],
    },
]


def build_document(spec: dict) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_pitch_letter(
        doc,
        spec["outlet"],
        spec["editor"],
        spec["open"],
        spec["angle"],
    )
    add_submission_instructions(doc, spec["outlet"], spec["instructions"])
    add_article_body(doc, spec["key"])

    path = OUT_DIR / spec["filename"]
    doc.save(path)
    print(f"Wrote {path}")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for spec in OUTLETS:
        build_document(spec)
    print(f"\nDone. {len(OUTLETS)} documents in:\n{OUT_DIR}")


if __name__ == "__main__":
    main()
