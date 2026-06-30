"""
Generate EIN Presswire Word document from the text press release.
Run from repo root: python marketing/press-releases/generate-ein-docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = Path(__file__).resolve().parent
TXT = HERE / "ein-miami-dade-aging-roof-index-2026.txt"
DOCX = HERE / "ein-miami-dade-aging-roof-index-2026.docx"


def main() -> None:
    if not TXT.exists():
        raise SystemExit(
            f"Missing {TXT.name}. Run: cd web && npx tsx scripts/generate-aging-roof-press-release.ts"
        )

    text = TXT.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        if block.startswith("FOR IMMEDIATE RELEASE"):
            p = doc.add_paragraph(block)
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if block.startswith("DataBloomer Publishes"):
            p = doc.add_paragraph(block)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            continue

        if block.startswith("MIAMI, Fla."):
            p = doc.add_paragraph(block)
            p.runs[0].italic = True
            continue

        if block.startswith("###"):
            doc.add_paragraph("###")
            continue

        if block.startswith("Top-ranked") or block.startswith("Key findings") or block.startswith("About DataBloomer") or block.startswith("Media Contact") or block.startswith("Note to editors"):
            doc.add_heading(block.split("\n")[0], level=2)
            rest = "\n".join(block.split("\n")[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
            continue

        if block.startswith("•"):
            for line in block.split("\n"):
                doc.add_paragraph(line.lstrip("• ").strip(), style="List Bullet")
            continue

        if block[0].isdigit() and ". ZIP" in block:
            for line in block.split("\n"):
                doc.add_paragraph(line, style="List Number")
            continue

        doc.add_paragraph(block)

    doc.save(DOCX)
    print(f"Wrote {DOCX}")


if __name__ == "__main__":
    main()
